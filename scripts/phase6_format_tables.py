#!/usr/bin/env python3
"""
Phase 6: Apply three-line table (三线表) formatting to all tables in the DOCX.

Formatting rules:
- Top border of table: 1 pt solid black
- Bottom border of header row (row 0): 0.5 pt solid black
- Bottom border of table (last row): 1 pt solid black
- NO left/right borders, NO vertical lines, NO other horizontal lines
- Table content font: 宋体 五号 (10.5pt) for Chinese, Times New Roman for Latin/numbers
- Header row bold

Usage:
    python scripts/phase6_format_tables.py
"""

from __future__ import annotations

import copy
import re
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn, nsdecls
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parents[1]
INPUT_DOCX = ROOT / "output" / "paper" / "论文完整版_phase6.docx"
OUTPUT_DOCX = ROOT / "output" / "paper" / "论文完整版_phase6.docx"  # overwrite in-place

# --- Border helpers ---

def make_border_element(tag: str, val: str = "single", sz: int = 8,
                        space: int = 0, color: str = "000000"):
    """Create a single border XML element.

    Args:
        tag: Border position (top, bottom, left, right, insideH, insideV, start, end)
        val: Border style - "single" for line, "none" / "nil" for no border
        sz: Size in eighths of a point (8 = 1pt, 4 = 0.5pt)
        space: Space in points
        color: Hex color string
    """
    el = OxmlElement(f"w:{tag}")
    el.set(qn("w:val"), val)
    el.set(qn("w:sz"), str(sz))
    el.set(qn("w:space"), str(space))
    el.set(qn("w:color"), color)
    return el


def clear_all_borders_table(tbl):
    """Set table-level borders to none/nil for all positions."""
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)

    # Remove existing tblBorders if present
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)

    borders = OxmlElement("w:tblBorders")
    for pos in ("top", "left", "bottom", "right", "insideH", "insideV"):
        borders.append(make_border_element(pos, val="none", sz=0))

    tblPr.append(borders)


def clear_cell_borders(cell):
    """Remove all borders from a single cell."""
    tcPr = cell._tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        cell._tc.insert(0, tcPr)

    existing = tcPr.find(qn("w:tcBorders"))
    if existing is not None:
        tcPr.remove(existing)

    borders = OxmlElement("w:tcBorders")
    for pos in ("top", "bottom", "left", "right"):
        borders.append(make_border_element(pos, val="nil", sz=0))
    tcPr.append(borders)


def set_cell_border(cell, position: str, sz: int = 8, val: str = "single",
                    color: str = "000000"):
    """Set a specific border on a cell.

    Args:
        cell: python-docx table cell
        position: 'top', 'bottom', 'left', 'right'
        sz: Size in eighths of a point
        val: 'single', 'none', etc.
        color: Hex color
    """
    tcPr = cell._tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        cell._tc.insert(0, tcPr)

    tc_borders = tcPr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tcPr.append(tc_borders)

    # Remove existing border for this position
    existing = tc_borders.find(qn(f"w:{position}"))
    if existing is not None:
        tc_borders.remove(existing)

    tc_borders.append(make_border_element(position, val=val, sz=sz, color=color))


def apply_three_line_table(table):
    """Apply three-line table formatting to a python-docx table.

    Three-line table rules:
    - Row 0 (header): top border 1pt, bottom border 0.5pt
    - Last row: bottom border 1pt
    - All other borders: none
    """
    n_rows = len(table.rows)
    if n_rows == 0:
        return

    # Step 1: Clear all table-level borders
    clear_all_borders_table(table._tbl)

    # Step 2: Clear all cell-level borders first
    for row in table.rows:
        for cell in row.cells:
            clear_cell_borders(cell)

    # Step 3: Apply the three lines

    # Line 1: Top of header row — 1pt (sz=8)
    for cell in table.rows[0].cells:
        set_cell_border(cell, "top", sz=8, val="single")

    # Line 2: Bottom of header row — 0.5pt (sz=4)
    for cell in table.rows[0].cells:
        set_cell_border(cell, "bottom", sz=4, val="single")

    # Line 3: Bottom of last row — 1pt (sz=8)
    for cell in table.rows[-1].cells:
        set_cell_border(cell, "bottom", sz=8, val="single")


# --- Font helpers ---

# Regex to separate CJK from Latin/digit characters
CJK_RANGE = re.compile(
    r"[\u4e00-\u9fff\u3400-\u4dbf\uf900-\ufaff"
    r"\U00020000-\U0002a6df\U0002a700-\U0002b73f"
    r"\U0002b740-\U0002b81f\U0002b820-\U0002ceaf"
    r"\U0002ceb0-\U0002ebef\U00030000-\U0003134f"
    r"\u3000-\u303f\uff00-\uffef]+"
)


def set_run_font(run, font_name_cjk: str = "宋体", font_name_latin: str = "Times New Roman",
                 size: Pt = Pt(10.5)):
    """Set font for a run, handling CJK and Latin separately."""
    run.font.size = size
    run.font.color.rgb = RGBColor(0, 0, 0)

    # Set the Latin font
    run.font.name = font_name_latin

    # Set the East Asian (CJK) font via XML
    rPr = run._r.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        run._r.insert(0, rPr)

    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)

    rFonts.set(qn("w:eastAsia"), font_name_cjk)
    rFonts.set(qn("w:ascii"), font_name_latin)
    rFonts.set(qn("w:hAnsi"), font_name_latin)


def format_table_fonts(table):
    """Set fonts for all cells in a table.

    Header row: bold, 宋体 五号 (10.5pt)
    Data rows:  normal, 宋体 五号 (10.5pt)
    Latin/numbers: Times New Roman
    """
    TABLE_FONT_SIZE = Pt(10.5)  # 五号 = 10.5pt

    for ri, row in enumerate(table.rows):
        is_header = (ri == 0)
        for cell in row.cells:
            for para in cell.paragraphs:
                # Center align all table cells
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # Set paragraph spacing to compact
                pf = para.paragraph_format
                pf.space_before = Pt(1)
                pf.space_after = Pt(1)

                for run in para.runs:
                    set_run_font(run, font_name_cjk="宋体",
                                 font_name_latin="Times New Roman",
                                 size=TABLE_FONT_SIZE)
                    run.bold = is_header


def set_table_width_autofit(table):
    """Set table to auto-fit to page width."""
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)

    # Set table width to 100%
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), "5000")
    tblW.set(qn("w:type"), "pct")

    # Set alignment to center
    jc = tblPr.find(qn("w:jc"))
    if jc is None:
        jc = OxmlElement("w:jc")
        tblPr.append(jc)
    jc.set(qn("w:val"), "center")


def remove_table_style(table):
    """Remove the default table style to prevent it from overriding borders."""
    tblPr = table._tbl.find(qn("w:tblPr"))
    if tblPr is not None:
        style_el = tblPr.find(qn("w:tblStyle"))
        if style_el is not None:
            tblPr.remove(style_el)


def format_caption_paragraph(para, is_table_caption: bool = True):
    """Format a table/figure caption paragraph.

    Table caption: centered, bold, 宋体五号 (10.5pt)
    Note paragraph: left-aligned, 小五号 (9pt), not bold
    """
    if is_table_caption:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = para.paragraph_format
        pf.space_before = Pt(6)
        pf.space_after = Pt(3)
        for run in para.runs:
            set_run_font(run, font_name_cjk="宋体",
                         font_name_latin="Times New Roman",
                         size=Pt(10.5))
            run.bold = True
    else:
        # Note paragraph (注：xxx)
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf = para.paragraph_format
        pf.space_before = Pt(2)
        pf.space_after = Pt(6)
        pf.first_line_indent = Pt(0)
        for run in para.runs:
            set_run_font(run, font_name_cjk="宋体",
                         font_name_latin="Times New Roman",
                         size=Pt(9))  # 小五号 = 9pt
            run.bold = False


def find_and_format_captions(doc):
    """Find and format table captions (above tables) and notes (below tables)."""
    import re
    caption_re = re.compile(r"^表\d+\s")
    note_re = re.compile(r"^注[：:]")

    body = doc.element.body
    elements = list(body)

    formatted_captions = 0
    formatted_notes = 0

    for i, el in enumerate(elements):
        if el.tag == qn("w:tbl"):
            # Check paragraph before table (caption)
            if i > 0 and elements[i - 1].tag == qn("w:p"):
                prev_para_el = elements[i - 1]
                prev_text = prev_para_el.text or ""
                if caption_re.match(prev_text.strip()):
                    # Find the matching Paragraph object
                    for para in doc.paragraphs:
                        if para._p is prev_para_el:
                            format_caption_paragraph(para, is_table_caption=True)
                            formatted_captions += 1
                            break

            # Check paragraph after table (note)
            if i < len(elements) - 1 and elements[i + 1].tag == qn("w:p"):
                next_para_el = elements[i + 1]
                next_text = next_para_el.text or ""
                if note_re.match(next_text.strip()):
                    for para in doc.paragraphs:
                        if para._p is next_para_el:
                            format_caption_paragraph(para, is_table_caption=False)
                            formatted_notes += 1
                            break

    return formatted_captions, formatted_notes


def main():
    print(f"Reading: {INPUT_DOCX}")
    doc = Document(str(INPUT_DOCX))
    n_tables = len(doc.tables)
    print(f"Found {n_tables} tables")

    for i, table in enumerate(doc.tables):
        n_rows = len(table.rows)
        n_cols = len(table.columns)
        header_text = table.rows[0].cells[0].text[:30] if n_rows > 0 else ""
        print(f"  Processing Table {i} ({n_rows}r x {n_cols}c) header='{header_text}'...")

        # Remove default table style that may override our borders
        remove_table_style(table)

        # Apply three-line table borders
        apply_three_line_table(table)

        # Format fonts
        format_table_fonts(table)

        # Set table width
        set_table_width_autofit(table)

    # Format captions and notes
    print("\nFormatting table captions and notes...")
    n_cap, n_note = find_and_format_captions(doc)
    print(f"  Formatted {n_cap} table captions (above tables)")
    print(f"  Formatted {n_note} note paragraphs (below tables)")

    doc.save(str(OUTPUT_DOCX))
    print(f"\nSaved: {OUTPUT_DOCX}")
    print("Done. All tables converted to three-line format (三线表).")


if __name__ == "__main__":
    main()
