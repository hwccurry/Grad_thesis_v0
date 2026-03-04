#!/usr/bin/env python3
"""
Post-process pandoc-generated .docx to match university thesis template.

Usage:
    python3 scripts/postprocess_docx.py

Input:  output/paper/论文终稿_base.docx  (pandoc output)
Output: output/paper/论文终稿.docx       (final formatted)
"""

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── paths ──────────────────────────────────────────────────────────
BASE = Path("/Users/mac/Desktop/Grad_thesis")
INPUT_DOCX = BASE / "output/paper/论文终稿_base.docx"
OUTPUT_DOCX = BASE / "output/paper/论文终稿.docx"

# ── font sizes (half-points) ──────────────────────────────────────
SZ_XIAO4 = 24    # 小四 = 12pt
SZ_5HAO  = 21    # 五号 = 10.5pt
SZ_XIAO5 = 18    # 小五 = 9pt

# ── helpers ────────────────────────────────────────────────────────

def set_run_font(run, ascii_font="Times New Roman", east_font="宋体",
                 size_half_pt=None, bold=None):
    """Set font properties on a run via XML for reliable CJK support."""
    rpr = run._element.find(qn('w:rPr'))
    if rpr is None:
        rpr = parse_xml(f'<w:rPr {nsdecls("w")}/>')
        run._element.insert(0, rpr)

    # fonts
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rpr.insert(0, rfonts)
    rfonts.set(qn('w:ascii'), ascii_font)
    rfonts.set(qn('w:hAnsi'), ascii_font)
    rfonts.set(qn('w:eastAsia'), east_font)
    rfonts.set(qn('w:cs'), ascii_font)

    # size
    if size_half_pt is not None:
        for tag in ('w:sz', 'w:szCs'):
            el = rpr.find(qn(tag))
            if el is None:
                el = parse_xml(f'<{tag} {nsdecls("w")}/>')
                rpr.append(el)
            el.set(qn('w:val'), str(size_half_pt))

    # bold
    if bold is not None:
        b_el = rpr.find(qn('w:b'))
        if bold:
            if b_el is None:
                rpr.append(parse_xml(f'<w:b {nsdecls("w")}/>'))
        else:
            if b_el is not None:
                rpr.remove(b_el)


def apply_three_line_borders(table):
    """Apply three-line table style: top 1pt, header-bottom 0.5pt, bottom 1pt."""
    # 1) Clear all cell borders first
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tc_pr = tc.tcPr
            if tc_pr is None:
                tc_pr = parse_xml(f'<w:tcPr {nsdecls("w")}/>')
                tc.insert(0, tc_pr)
            existing = tc_pr.find(qn('w:tcBorders'))
            if existing is not None:
                tc_pr.remove(existing)

    # 2) Set borders per row
    num_rows = len(table.rows)
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            tc = cell._tc
            tc_pr = tc.tcPr
            if tc_pr is None:
                tc_pr = parse_xml(f'<w:tcPr {nsdecls("w")}/>')
                tc.insert(0, tc_pr)

            borders_xml = f'<w:tcBorders {nsdecls("w")}>'

            # Top border: only first row gets 1pt top
            if row_idx == 0:
                borders_xml += '<w:top w:val="single" w:sz="8" w:space="0" w:color="000000"/>'
            else:
                borders_xml += '<w:top w:val="nil"/>'

            # Bottom border
            if row_idx == 0 and num_rows == 1:
                # Single-row table: 1pt bottom
                borders_xml += '<w:bottom w:val="single" w:sz="8" w:space="0" w:color="000000"/>'
            elif row_idx == 0:
                # Header row bottom: 0.5pt
                borders_xml += '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            elif row_idx == num_rows - 1:
                # Last row bottom: 1pt
                borders_xml += '<w:bottom w:val="single" w:sz="8" w:space="0" w:color="000000"/>'
            else:
                borders_xml += '<w:bottom w:val="nil"/>'

            # No left/right borders, no internal lines
            borders_xml += '<w:left w:val="nil"/>'
            borders_xml += '<w:right w:val="nil"/>'
            borders_xml += '<w:insideH w:val="nil"/>'
            borders_xml += '<w:insideV w:val="nil"/>'
            borders_xml += '</w:tcBorders>'
            tc_pr.append(parse_xml(borders_xml))


def set_table_cell_fonts(table, ascii_font="Times New Roman", east_font="宋体",
                         size_half_pt=SZ_5HAO):
    """Set font for all runs in all cells of a table."""
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(0)
                para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.paragraph_format.first_line_indent = None
                for run in para.runs:
                    set_run_font(run, ascii_font, east_font, size_half_pt)


def set_table_header_bold(table):
    """Make first row bold."""
    if len(table.rows) > 0:
        for cell in table.rows[0].cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    set_run_font(run, bold=True)
                    rpr = run._element.find(qn('w:rPr'))
                    if rpr is not None:
                        b = rpr.find(qn('w:b'))
                        if b is None:
                            rpr.append(parse_xml(f'<w:b {nsdecls("w")}/>'))


def _first_nonempty_run_is_bold(para):
    """Return True only when first non-empty run is explicitly bold."""
    for run in para.runs:
        if run.text and run.text.strip():
            return run.bold is True
    return False


def is_table_caption(text, para=None):
    """Check if paragraph is a table caption (表X title or 附表AX title).
    Requires: number followed by space, paragraph short, first run bold."""
    t = text.strip()
    if not re.match(r'^(表\d+|附表A?\d+)[ \u3000]+', t):
        return False
    if len(t) > 80:
        return False
    if para:
        return _first_nonempty_run_is_bold(para)
    return False


def is_figure_caption(text, para=None):
    """Check if paragraph is a figure caption (图X title or 附图AX title)."""
    t = text.strip()
    if not re.match(r'^(图\d+|附图A?\d+)[ \u3000]+', t):
        return False
    if len(t) > 80:
        return False
    if para:
        return _first_nonempty_run_is_bold(para)
    return False


def is_note(text):
    """Check if paragraph is a note line."""
    t = text.strip()
    return t.startswith('注：') or t.startswith('注:')


def has_image(para):
    """Check if paragraph contains an inline image."""
    xml = para._element.xml
    return 'w:drawing' in xml and 'a:blip' in xml


def has_display_omml(para):
    """True for standalone display equation paragraphs (OMML math)."""
    xml = para._element.xml
    return 'm:oMathPara' in xml or ('m:oMath' in xml and not para.text.strip())


def add_page_numbers(doc):
    """Add centered page numbers to footer."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        # Clear existing footer content
        for p in footer.paragraphs:
            p.clear()
        # Get or create footer paragraph
        if len(footer.paragraphs) == 0:
            fp = footer.add_paragraph()
        else:
            fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add PAGE field
        run = fp.add_run()
        set_run_font(run, "Times New Roman", "宋体", SZ_XIAO5)
        run._element.append(
            parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'))

        run2 = fp.add_run()
        set_run_font(run2, "Times New Roman", "宋体", SZ_XIAO5)
        run2._element.append(
            parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>'))

        run3 = fp.add_run()
        set_run_font(run3, "Times New Roman", "宋体", SZ_XIAO5)
        run3._element.append(
            parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'))


def resize_images(doc, max_width_cm=14.0):
    """Resize all inline images to fit within max_width, preserving aspect ratio."""
    max_width_emu = Cm(max_width_cm)
    body = doc.element.body
    for drawing in body.iter(qn('w:drawing')):
        for inline in drawing.iter(qn('wp:inline')):
            extent = inline.find(qn('wp:extent'))
            if extent is not None:
                cx = int(extent.get('cx', 0))
                cy = int(extent.get('cy', 0))
                if cx > max_width_emu and cx > 0:
                    ratio = max_width_emu / cx
                    new_cx = int(cx * ratio)
                    new_cy = int(cy * ratio)
                    extent.set('cx', str(new_cx))
                    extent.set('cy', str(new_cy))
                    for ext_el in inline.iter(qn('a:ext')):
                        ecx = int(ext_el.get('cx', 0))
                        if ecx > max_width_emu:
                            ext_el.set('cx', str(new_cx))
                            ext_el.set('cy', str(new_cy))


# ── main processing ───────────────────────────────────────────────

def main():
    print(f"Loading {INPUT_DOCX} ...")
    doc = Document(str(INPUT_DOCX))

    # ── 1. Page setup ──────────────────────────────────────────────
    print("Setting page margins and layout ...")
    for section in doc.sections:
        section.page_width = Cm(21.0)   # A4
        section.page_height = Cm(29.7)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(3.0)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)

    # ── 2. Process paragraphs ──────────────────────────────────────
    print("Processing paragraphs ...")
    paragraphs = doc.paragraphs

    in_references = False  # track if we're inside the references section
    in_declaration = False  # track if we're inside the declaration section

    for i, para in enumerate(paragraphs):
        text = para.text.strip()
        style_name = para.style.name

        # Track section boundaries
        if style_name == 'Heading 1':
            in_references = (text == '参考文献')
            in_declaration = (text == '论文独创性及授权声明')

        # ── Display math (OMML) → centered, no indent ──
        if has_display_omml(para):
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.first_line_indent = Cm(0)
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(6)
            continue

        # ── 2a. Heading styles ──
        if style_name.startswith('Heading'):
            if style_name == 'Heading 1':
                para.paragraph_format.first_line_indent = Cm(0)
                para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif style_name == 'Heading 2':
                para.paragraph_format.first_line_indent = Cm(0)
                para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            elif style_name == 'Heading 3':
                para.paragraph_format.first_line_indent = Cm(0)
                para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            continue

        # ── 2a-1. References → hanging indent [1] format ──
        if in_references and re.match(r'^\[\d+\]\s*', text):
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.left_indent = Pt(24)
            para.paragraph_format.first_line_indent = Pt(-24)
            para.paragraph_format.line_spacing = 1.5
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            for run in para.runs:
                set_run_font(run, "Times New Roman", "宋体", SZ_XIAO4)
            continue

        # ── 2a-2. Declaration signature line → no indent, right-aligned ──
        if in_declaration and ('作者签名' in text or '导师签名' in text):
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.first_line_indent = Cm(0)
            para.paragraph_format.space_before = Pt(24)
            for run in para.runs:
                set_run_font(run, "Times New Roman", "宋体", SZ_XIAO4)
            continue

        # ── 2b. Table captions → centered, bold, 宋体五号 ──
        if is_table_caption(text, para):
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.first_line_indent = Cm(0)
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(3)
            for run in para.runs:
                set_run_font(run, "Times New Roman", "宋体", SZ_5HAO, bold=True)
                rpr = run._element.find(qn('w:rPr'))
                if rpr is not None and rpr.find(qn('w:b')) is None:
                    rpr.append(parse_xml(f'<w:b {nsdecls("w")}/>'))
            continue

        # ── 2c. Figure captions → centered, bold, 宋体五号 ──
        if is_figure_caption(text, para):
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.first_line_indent = Cm(0)
            para.paragraph_format.space_before = Pt(3)
            para.paragraph_format.space_after = Pt(6)
            for run in para.runs:
                set_run_font(run, "Times New Roman", "宋体", SZ_5HAO, bold=True)
                rpr = run._element.find(qn('w:rPr'))
                if rpr is not None and rpr.find(qn('w:b')) is None:
                    rpr.append(parse_xml(f'<w:b {nsdecls("w")}/>'))
            continue

        # ── 2d. Notes → 小五号(9pt), left-aligned ──
        if is_note(text):
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.first_line_indent = Cm(0)
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after = Pt(6)
            for run in para.runs:
                set_run_font(run, "Times New Roman", "宋体", SZ_XIAO5)
            continue

        # ── 2e. Image paragraphs → center ──
        if has_image(para) and not text:
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.first_line_indent = Cm(0)
            continue

        # ── 2f. Keywords line → no indent ──
        if text.startswith('关键词') or text.startswith('Keywords'):
            para.paragraph_format.first_line_indent = Cm(0)
            para.paragraph_format.space_before = Pt(6)
            for run in para.runs:
                set_run_font(run, "Times New Roman", "宋体", SZ_XIAO4)
            continue

        # ── 2g. Normal body text → 宋体小四, first-line indent, 1.5x spacing ──
        if style_name == 'Normal':
            para.paragraph_format.first_line_indent = Pt(24)  # ~2 chars at 12pt
            para.paragraph_format.line_spacing = 1.5
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            for run in para.runs:
                set_run_font(run, "Times New Roman", "宋体", SZ_XIAO4)

    # ── 3. Process tables ──────────────────────────────────────────
    print(f"Processing {len(doc.tables)} tables ...")
    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        apply_three_line_borders(table)
        set_table_cell_fonts(table)
        set_table_header_bold(table)

    # ── 4. Resize images ──────────────────────────────────────────
    print("Resizing images ...")
    resize_images(doc, max_width_cm=14.0)

    # ── 5. Add page numbers ──────────────────────────────────────
    print("Adding page numbers ...")
    add_page_numbers(doc)

    # ── 6. Save ───────────────────────────────────────────────────
    print(f"Saving to {OUTPUT_DOCX} ...")
    doc.save(str(OUTPUT_DOCX))
    print("Done!")


if __name__ == "__main__":
    main()
