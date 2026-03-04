#!/usr/bin/env python3
"""
Phase 6 delivery audit.

Checks:
1) Placeholder scan in full markdown draft.
2) Figure link existence + DPI check (>=300).
3) Figure/Table caption sequence check.
4) DID reproducibility consistency check.
5) DOCX export existence and structure summary.

Output:
  notes/phase6_delivery_audit.md
"""

from __future__ import annotations

from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
import re

import pandas as pd
from PIL import Image
from docx import Document


ROOT = Path(__file__).resolve().parents[1]
PAPER_MD = ROOT / "output" / "paper" / "论文完整版.md"
PAPER_DOCX = ROOT / "output" / "paper" / "论文终稿.docx"
BASELINE_CSV = ROOT / "output" / "tables" / "did_baseline_regression.csv"
REPRO_CSV = ROOT / "output" / "tables" / "phase5_did_repro_check.csv"
REPORT = ROOT / "notes" / "phase6_delivery_audit.md"

PLACEHOLDER_PATTERN = re.compile(r"\bR\d{2}\b|TODO|TBD|待补|待定|xxx|XXX")
IMAGE_PATTERN = re.compile(r"!\[[^\]]*]\(([^)]+)\)")
FIG_CAPTION_PATTERN = re.compile(r"\*\*图(\d+)\s")
TAB_CAPTION_PATTERN = re.compile(r"\*\*表(\d+)\s")
REF_ENTRY_PATTERN = re.compile(r"^\[\d+\]\s", re.MULTILINE)

# Required sections that must exist as H1 headings in the merged paper
REQUIRED_SECTIONS = [
    "摘要",
    "Abstract",
    "第一章 绪论",
    "第二章 文献综述与理论基础",
    "第三章 上市公司分红的动因预测分析",
    "第四章 基于政策准自然实验的因果效应评估",
    "第五章 研究结论、启示与展望",
    "参考文献",
    "后记",
    "论文独创性及授权声明",
]
MIN_REF_COUNT = 25
DPI_MIN = 299.5


@dataclass
class DidRow:
    outcome: str
    controls: str
    coef: float
    tval: float


def now_str() -> str:
    return datetime.now().astimezone().strftime("%Y-%m-%d %H:%M:%S %z")


def is_sequential(nums: list[int]) -> tuple[bool, str]:
    if not nums:
        return True, "无编号"
    expected = list(range(nums[0], nums[0] + len(nums)))
    ok = nums == expected
    if ok:
        return True, f"{nums[0]}–{nums[-1]} 连续"
    return False, f"实际={nums}, 期望={expected}"


def load_text(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def check_placeholders(text: str) -> list[str]:
    return PLACEHOLDER_PATTERN.findall(text)


def check_images(text: str) -> tuple[list[dict], bool]:
    rows: list[dict] = []
    all_ok = True
    md_dir = PAPER_MD.parent
    links = IMAGE_PATTERN.findall(text)
    for link in links:
        img_path = (md_dir / link).resolve()
        exists = img_path.exists()
        dpi_ok = False
        dpi_text = "N/A"
        if exists:
            try:
                with Image.open(img_path) as img:
                    dpi = img.info.get("dpi")
                    if isinstance(dpi, tuple) and len(dpi) >= 2:
                        xdpi, ydpi = float(dpi[0]), float(dpi[1])
                        dpi_text = f"{xdpi:.1f}x{ydpi:.1f}"
                        dpi_ok = xdpi >= DPI_MIN and ydpi >= DPI_MIN
                    else:
                        dpi_text = "missing"
                        dpi_ok = False
            except Exception as exc:
                dpi_text = f"error:{type(exc).__name__}"
                dpi_ok = False
        status = exists and dpi_ok
        all_ok = all_ok and status
        rows.append(
            {
                "link": link,
                "resolved": str(img_path.relative_to(ROOT)) if exists else str(img_path),
                "exists": exists,
                "dpi": dpi_text,
                "ok": status,
            }
        )
    return rows, all_ok


def load_did_baseline(path: Path) -> list[DidRow]:
    df = pd.read_csv(path)
    rows = []
    for _, r in df.iterrows():
        rows.append(
            DidRow(
                outcome=str(r["被解释变量"]),
                controls=str(r["控制变量"]),
                coef=float(r["did系数"]),
                tval=float(r["t值"]),
            )
        )
    return rows


def load_did_repro(path: Path) -> list[DidRow]:
    df = pd.read_csv(path)
    rows = []
    for _, r in df.iterrows():
        controls = "是" if str(r["spec"]) == "with_controls" else "否"
        rows.append(
            DidRow(
                outcome=str(r["outcome"]),
                controls=controls,
                coef=float(r["did_coef"]),
                tval=float(r["did_t"]),
            )
        )
    return rows


def compare_did(baseline: list[DidRow], repro: list[DidRow]) -> tuple[bool, list[str]]:
    idx = {(r.outcome, r.controls): r for r in baseline}
    details = []
    ok = True
    for rr in repro:
        key = (rr.outcome, rr.controls)
        br = idx.get(key)
        if br is None:
            ok = False
            details.append(f"{rr.outcome}/{rr.controls}: baseline 缺失")
            continue
        d_coef = abs(rr.coef - br.coef)
        d_t = abs(rr.tval - br.tval)
        same = d_coef <= 5e-4 and d_t <= 5e-3
        ok = ok and same
        details.append(
            f"{rr.outcome}/{rr.controls}: Δcoef={d_coef:.6f}, Δt={d_t:.6f}, {'ok' if same else 'mismatch'}"
        )
    return ok, details


def docx_summary(path: Path) -> dict:
    if not path.exists():
        return {"exists": False}
    doc = Document(str(path))
    image_rels = sum(1 for r in doc.part._rels.values() if "image" in r.reltype)
    heading_count = sum(
        1
        for p in doc.paragraphs
        if p.style is not None and str(p.style.name).startswith("Heading")
    )
    return {
        "exists": True,
        "size_bytes": path.stat().st_size,
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "images": image_rels,
        "heading_paras": heading_count,
    }


def check_sections(text: str) -> tuple[bool, list[str]]:
    """Check that all required sections exist as H1 headings."""
    h1_pattern = re.compile(r"^# (.+)$", re.MULTILINE)
    h1_titles = [m.group(1).strip() for m in h1_pattern.finditer(text)]
    missing = []
    for req in REQUIRED_SECTIONS:
        if not any(req in t for t in h1_titles):
            missing.append(req)
    return len(missing) == 0, missing


def check_ref_count(text: str) -> tuple[bool, int]:
    """Check that reference count meets minimum threshold."""
    count = len(REF_ENTRY_PATTERN.findall(text))
    return count >= MIN_REF_COUNT, count


def build_report() -> str:
    text = load_text(PAPER_MD)
    placeholder_hits = check_placeholders(text)

    image_rows, image_ok = check_images(text)
    fig_nums = [int(x) for x in FIG_CAPTION_PATTERN.findall(text)]
    tab_nums = [int(x) for x in TAB_CAPTION_PATTERN.findall(text)]
    fig_seq_ok, fig_seq_msg = is_sequential(fig_nums)
    tab_seq_ok, tab_seq_msg = is_sequential(tab_nums)

    baseline = load_did_baseline(BASELINE_CSV)
    repro = load_did_repro(REPRO_CSV)
    did_ok, did_details = compare_did(baseline, repro)

    docx = docx_summary(PAPER_DOCX)

    sections_ok, sections_missing = check_sections(text)
    ref_ok, ref_count = check_ref_count(text)

    lines = [
        "# Phase 6 交付审计报告",
        "",
        f"- 生成时间：{now_str()}",
        f"- 审计对象：`{PAPER_MD.relative_to(ROOT)}`、`{PAPER_DOCX.relative_to(ROOT)}`",
        "",
        "## 1. 文本占位符检查",
        f"- 结果：{'通过' if len(placeholder_hits) == 0 else '未通过'}",
        f"- 命中数量：{len(placeholder_hits)}",
    ]
    if placeholder_hits:
        lines.append(f"- 命中明细：{', '.join(placeholder_hits)}")
    else:
        lines.append("- 命中明细：无")

    lines += [
        "",
        "## 2. 图像链接与分辨率检查",
        f"- 图像总数：{len(image_rows)}",
        f"- 结果：{'通过' if image_ok else '未通过'}",
        "",
        "| 链接 | 文件 | 存在 | DPI | 通过 |",
        "| --- | --- | --- | --- | --- |",
    ]
    for r in image_rows:
        lines.append(
            f"| `{r['link']}` | `{r['resolved']}` | {'是' if r['exists'] else '否'} | {r['dpi']} | {'是' if r['ok'] else '否'} |"
        )

    lines += [
        "",
        "## 3. 表图编号连续性",
        f"- 图编号检查：{'通过' if fig_seq_ok else '未通过'}（{fig_seq_msg}）",
        f"- 表编号检查：{'通过' if tab_seq_ok else '未通过'}（{tab_seq_msg}）",
        "",
        "## 4. DID 关键系数复验",
        f"- 结果：{'通过' if did_ok else '未通过'}",
    ]
    for d in did_details:
        lines.append(f"- {d}")

    lines += [
        "",
        "## 5. 章节完整性门禁",
        f"- 结果：{'通过' if sections_ok else '未通过'}",
    ]
    if sections_missing:
        lines.append(f"- 缺失章节：{', '.join(sections_missing)}")
    else:
        lines.append("- 全部必需章节均已包含")

    lines += [
        "",
        "## 6. 参考文献数量检查",
        f"- 结果：{'通过' if ref_ok else '未通过'}",
        f"- 参考文献条目数：{ref_count}（下限 {MIN_REF_COUNT}）",
    ]

    lines += [
        "",
        "## 7. DOCX 导出结构摘要",
    ]
    if not docx.get("exists"):
        lines.append("- 结果：未导出 docx")
    else:
        lines.append("- 结果：已导出")
        lines.append(f"- 文件大小：{docx['size_bytes']} bytes")
        lines.append(f"- 段落数：{docx['paragraphs']}")
        lines.append(f"- 表格数：{docx['tables']}")
        lines.append(f"- 图片数：{docx['images']}")
        lines.append(f"- Heading 段落数：{docx['heading_paras']}")

    overall_ok = (
        len(placeholder_hits) == 0
        and image_ok
        and fig_seq_ok
        and tab_seq_ok
        and did_ok
        and sections_ok
        and ref_ok
        and docx.get("exists", False)
    )
    lines += [
        "",
        "## 8. 结论",
        f"- 自动化审计结论：{'通过' if overall_ok else '存在未通过项'}",
        "- 说明：Word 版式细节（页眉页脚、三线表线宽、字体混排）仍需人工终检。",
        "",
    ]

    return "\n".join(lines)


def main() -> None:
    report = build_report()
    REPORT.parent.mkdir(parents=True, exist_ok=True)
    REPORT.write_text(report, encoding="utf-8")
    print(f"saved: {REPORT}")


if __name__ == "__main__":
    main()
